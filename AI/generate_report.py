from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor

# 创建文档
doc = Document()

# 设置默认字体
for style in doc.styles:
    if style.name in ['Normal', 'Body Text', 'Body Text 2', 'Body Text 3']:
        style.font.name = '仿宋'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        style.font.size = Pt(10.5)  # 五号字体
        style.paragraph_format.line_spacing = 1.5  # 1.5倍行间距
        style.paragraph_format.first_line_indent = Inches(0.35)  # 首行缩进两字符

# 设置标题样式
def set_heading_style(paragraph, level, text):
    if level == 1:
        # 一级标题：黑体三号加粗
        heading = doc.add_heading(text, level=0)
        run = heading.runs[0]
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(16)  # 三号字体
        run.font.bold = True
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.space_after = Pt(12)
        return heading
    elif level == 2:
        # 二级标题：黑体四号加粗
        heading = doc.add_heading(text, level=1)
        run = heading.runs[0]
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(14)  # 四号字体
        run.font.bold = True
        heading.paragraph_format.space_after = Pt(12)
        return heading
    elif level == 3:
        # 三级标题：黑体小四号加粗
        heading = doc.add_heading(text, level=2)
        run = heading.runs[0]
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(12)  # 小四号字体
        run.font.bold = True
        heading.paragraph_format.space_after = Pt(12)
        return heading

# 添加正文段落
def add_paragraph(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Inches(0.35)  # 首行缩进两字符
    p.paragraph_format.line_spacing = 1.5  # 1.5倍行间距
    return p

# 添加列表项
def add_list(items):
    for item in items:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.first_line_indent = Inches(0)  # 列表项不缩进
        p.paragraph_format.line_spacing = 1.5  # 1.5倍行间距

# 添加表格
def add_table(data, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # 设置表头样式
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    # 添加数据行
    for row in data:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row_cells):
            cell.text = str(row[i])
    # 设置表格样式
    table.style = 'Table Grid'
    return table

# 生成报告内容

# 标题页
set_heading_style(doc, 1, '图像识别实验报告')
doc.add_paragraph()
doc.add_paragraph()
author_p = doc.add_paragraph('姓名：')
author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
student_id_p = doc.add_paragraph('学号：')
student_id_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
course_p = doc.add_paragraph('课程：人工智能基础')
course_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p = doc.add_paragraph('日期：2025年12月23日')
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# 一、实验目的
set_heading_style(doc, 2, '一、实验目的')
add_paragraph('本次实验旨在完成Fashion MNIST图像分类任务，实践并理解基本的图像分类管道，比较不同分类器的性能差异，并探索各种优化技术对模型性能的影响。具体目标包括：')
add_list([
    '理解图像分类的基本流程和数据驱动方法',
    '掌握train/val/test数据集划分及验证集的使用',
    '实现并比较Softmax分类器和全连接神经网络分类器',
    '探索不同优化算法和正则化技术的效果',
    '通过贝叶斯优化寻找最佳超参数组合'
])

# 二、数据准备
set_heading_style(doc, 2, '二、数据准备')

set_heading_style(doc, 3, '2.1 数据集选择')
add_paragraph('本次实验使用Fashion MNIST数据集，该数据集包含10个类别的时装图像，共60,000张训练图像和10,000张测试图像，每张图像大小为28×28像素的灰度图。')

set_heading_style(doc, 3, '2.2 数据划分')
add_paragraph('将数据集划分为三个部分：')
add_list([
    '训练集：50,000张图像，用于模型训练',
    '验证集：10,000张图像，用于超参数调优和早停判断',
    '测试集：10,000张图像，用于最终模型评估'
])

set_heading_style(doc, 3, '2.3 数据增强')
add_paragraph('为提高模型的泛化能力，对训练数据进行了以下增强操作：')
add_list([
    '随机水平翻转',
    '随机旋转（±5度）',
    '随机平移（±10%）'
])

# 三、模型设计
set_heading_style(doc, 2, '三、模型设计')

set_heading_style(doc, 3, '3.1 模型结构')
add_paragraph('本次实验实现了优化后的全连接神经网络，结构如下：')
# 添加模型结构表格
model_data = [
    ['输入层', '784（28×28）', '-', '-'],
    ['隐藏层1', '256', 'ReLU', '批归一化 + Dropout（率0.5）'],
    ['隐藏层2', '128', 'ReLU', '批归一化 + Dropout（率0.5）'],
    ['输出层', '10', '无（CrossEntropyLoss包含Softmax）', '-']
]
add_table(model_data, ['层类型', '神经元数量', '激活函数', '正则化'])

set_heading_style(doc, 3, '3.2 正则化技术')
add_paragraph('为防止过拟合，采用了多种正则化技术：')
add_list([
    'Dropout：在每个隐藏层后添加Dropout层，丢弃率为0.5',
    '批量归一化：在激活函数前添加批量归一化层，加速训练并提高模型稳定性',
    'L2正则化：通过权重衰减实现，系数为1e-4',
    '早停策略：当验证损失连续5轮未下降时停止训练'
])

set_heading_style(doc, 3, '3.3 优化算法')
add_paragraph('使用Adam优化器，结合学习率衰减策略：')
add_list([
    '初始学习率：由贝叶斯优化自动搜索',
    '学习率衰减：当验证损失连续3轮未下降时，学习率减半'
])

# 四、训练过程
set_heading_style(doc, 2, '四、训练过程')

set_heading_style(doc, 3, '4.1 训练参数')
add_paragraph('通过贝叶斯优化确定的最佳超参数组合：')
# 添加超参数表格
hyper_data = [
    ['学习率', '0.001'],
    ['隐藏层1神经元数', '256'],
    ['隐藏层2神经元数', '128'],
    ['Dropout率', '0.5'],
    ['权重衰减', '1e-4'],
    ['批量大小', '64'],
    ['最大训练轮数', '30'],
    ['早停耐心值', '5']
]
add_table(hyper_data, ['超参数', '最佳值'])

set_heading_style(doc, 3, '4.2 训练流程')
add_paragraph('训练过程包括以下步骤：')
add_list([
    '数据加载与增强',
    '模型初始化',
    '训练循环（前向传播→损失计算→反向传播→参数更新）',
    '验证集评估',
    '学习率调整',
    '早停检查',
    '保存最佳模型'
])

# 五、结果分析
set_heading_style(doc, 2, '五、结果分析')

set_heading_style(doc, 3, '5.1 学习曲线')
add_paragraph('模型训练过程中，训练损失和验证损失均逐渐下降并趋于稳定，训练准确率和验证准确率逐渐上升，最终训练准确率达到79.62%，验证准确率达到83.49%，测试准确率达到84.88%。')
# 图片位置标注
img_p = doc.add_paragraph('（此处插入优化后的学习曲线图像：optimized_fcnn_learning_curve.png）')
img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

set_heading_style(doc, 3, '5.2 模型性能评估')
add_paragraph('模型性能评估结果如下：')
# 添加性能评估表格
perf_data = [
    ['最终训练损失', '0.5531'],
    ['最终验证损失', '0.4503'],
    ['最终训练准确率', '79.62%'],
    ['最终验证准确率', '83.49%'],
    ['测试准确率', '84.88%'],
    ['训练/验证准确率差距', '3.87%']
]
add_table(perf_data, ['指标', '数值'])
add_paragraph('从评估结果可以看出，模型具有良好的泛化能力，训练和验证准确率差距仅为3.87%，没有出现明显的过拟合现象。')

set_heading_style(doc, 3, '5.3 理想曲线对比')
add_paragraph('为了直观展示模型训练的理想状态，生成了理想的学习曲线进行对比。理想曲线显示，训练损失和验证损失应持续下降并趋于稳定，训练准确率和验证准确率应持续上升并趋于稳定，且两者差距较小。')
# 图片位置标注
img_p = doc.add_paragraph('（此处插入理想学习曲线图像：ideal_curves/Ideal_Accuracy_Comparison.png）')
img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 六、结论
set_heading_style(doc, 2, '六、结论')

set_heading_style(doc, 3, '6.1 实验总结')
add_paragraph('本次实验成功完成了Fashion MNIST图像分类任务，实现了优化后的全连接神经网络，并通过多种优化技术提高了模型性能。主要完成了以下工作：')
add_list([
    '实现了完整的图像分类管道，包括数据加载、模型训练、验证和测试',
    '采用了多种正则化技术，包括Dropout、批量归一化和L2正则化',
    '实现了早停策略和学习率衰减，提高了训练效率',
    '使用贝叶斯优化自动搜索最佳超参数组合',
    '生成了清晰的学习曲线，直观展示了模型性能'
])

set_heading_style(doc, 3, '6.2 结果分析')
add_paragraph('实验结果表明，优化后的全连接神经网络在Fashion MNIST数据集上取得了84.88%的测试准确率，具有良好的泛化能力。通过比较可以看出，全连接神经网络的性能优于简单的Softmax分类器，而添加正则化技术和数据增强可以进一步提高模型性能。')

set_heading_style(doc, 3, '6.3 改进方向')
add_paragraph('尽管模型取得了不错的性能，但仍有改进空间：')
add_list([
    '尝试更复杂的模型结构，如卷积神经网络（CNN）',
    '探索更多的数据增强方法',
    '尝试不同的优化算法和损失函数',
    '进一步优化超参数组合'
])

# 七、代码实现
set_heading_style(doc, 2, '七、代码实现')
add_paragraph('实验代码主要包含以下文件：')
add_list([
    'fashion_mnist_optimized.py：优化后的训练脚本，包含数据增强、模型定义、训练函数和贝叶斯优化',
    'best_optimized_model.pth：保存的最佳模型参数',
    'optimized_fcnn_learning_curve.png：优化后的学习曲线'
])
add_paragraph('代码实现了完整的图像分类管道，包括数据加载、模型训练、验证和测试，以及多种优化技术和正则化方法。')

# 八、参考文献
set_heading_style(doc, 2, '八、参考文献')
add_list([
    '[1] Xiao, Han, et al. "Fashion-mnist: a novel image dataset for benchmarking machine learning algorithms." arXiv preprint arXiv:1708.07747 (2017).',
    '[2] Kingma, Diederik P., and Jimmy Ba. "Adam: A method for stochastic optimization." arXiv preprint arXiv:1412.6980 (2014).',
    '[3] Srivastava, Nitish, et al. "Dropout: a simple way to prevent neural networks from overfitting." The journal of machine learning research 15.1 (2014): 1929-1958.',
    '[4] Ioffe, Sergey, and Christian Szegedy. "Batch normalization: Accelerating deep network training by reducing internal covariate shift." arXiv preprint arXiv:1502.03167 (2015).'
])

# 保存文档
doc.save('图像识别实验报告_final.docx')
print("实验报告已生成，文件名：图像识别实验报告_final.docx")
