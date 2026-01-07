#!/usr/bin/env python3
# 将Markdown用户手册转换为DOCX文档

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

class MarkdownToDocxConverter:
    def __init__(self):
        self.doc = Document()
        # 设置默认样式
        self._setup_default_styles()
        
    def _setup_default_styles(self):
        """设置默认样式"""
        # 设置正文样式
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '仿宋'
        font.size = Pt(12)  # 小四对应12磅
        font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        
        # 设置段落格式
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.0  # 单倍行距
        paragraph_format.first_line_indent = Pt(24)  # 首行缩进两字符，约24磅
        
    def _add_paragraph(self, text, style=None, is_title=False):
        """添加段落，根据是否为标题设置不同样式"""
        paragraph = self.doc.add_paragraph(text)
        
        if is_title:
            # 标题使用正文格式，不使用标题样式
            font = paragraph.runs[0].font
            font.name = '黑体'
            font.bold = True
            font.color.rgb = RGBColor(0, 0, 0)  # 黑色
            paragraph.paragraph_format.first_line_indent = Pt(0)  # 标题不缩进
        else:
            # 正文样式
            for run in paragraph.runs:
                run.font.name = '仿宋'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
            paragraph.paragraph_format.first_line_indent = Pt(24)  # 首行缩进两字符
        
        return paragraph
    
    def _add_heading(self, text, level):
        """添加标题，根据级别设置不同大小"""
        paragraph = self.doc.add_paragraph(text)
        font = paragraph.runs[0].font
        font.name = '黑体'
        font.bold = True
        font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        
        # 根据级别设置字体大小
        if level == 1:
            font.size = Pt(16)  # 三号对应16磅
        elif level == 2:
            font.size = Pt(14)  # 四号对应14磅
        elif level == 3:
            font.size = Pt(12)  # 小四对应12磅
        elif level == 4:
            font.size = Pt(12)  # 小四对应12磅
        
        # 标题不缩进
        paragraph.paragraph_format.first_line_indent = Pt(0)
        
        return paragraph
    
    def _add_table(self, markdown_table):
        """添加表格"""
        lines = markdown_table.strip().split('\n')
        if not lines:
            return
        
        # 解析表格
        header = None
        rows = []
        
        for line in lines:
            if line.strip().startswith('|'):
                # 移除前后的|，然后分割
                cells = [cell.strip() for cell in line.strip()[1:-1].split('|')]
                if not header:
                    header = cells
                else:
                    # 跳过分隔线
                    if all(cell == '-' * len(cell) for cell in cells):
                        continue
                    rows.append(cells)
        
        if header and rows:
            # 创建表格
            table = self.doc.add_table(rows=1, cols=len(header), style='Table Grid')
            
            # 设置表头
            hdr_cells = table.rows[0].cells
            for i, cell in enumerate(header):
                hdr_cells[i].text = cell
                # 设置表头样式
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '仿宋'
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
                
            # 添加数据行
            for row_data in rows:
                row_cells = table.add_row().cells
                for i, cell in enumerate(row_data):
                    row_cells[i].text = cell
                    # 设置单元格样式
                    for paragraph in row_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '仿宋'
                            run.font.size = Pt(12)
                            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
    
    def convert(self, markdown_file, output_file):
        """将Markdown文件转换为DOCX"""
        with open(markdown_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 分割Markdown内容
        lines = content.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # 检查是否为标题
            if line.startswith('#'):
                # 计算标题级别
                level = len(line.split(' ')[0])
                title_text = line[level:].strip()
                self._add_heading(title_text, level)
                i += 1
            # 检查是否为表格
            elif line.startswith('|'):
                # 收集整个表格
                table_lines = [line]
                i += 1
                while i < len(lines) and (lines[i].strip().startswith('|') or lines[i].strip().startswith('+')):
                    table_lines.append(lines[i])
                    i += 1
                table_content = '\n'.join(table_lines)
                self._add_table(table_content)
            # 检查是否为代码块
            elif line.startswith('```'):
                # 处理代码块
                code_content = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_content.append(lines[i])
                    i += 1
                i += 1
                
                # 将代码块内容作为普通段落添加
                if code_content:
                    code_text = '\n'.join(code_content)
                    self._add_paragraph(code_text, is_title=False)
            # 检查是否为空行
            elif not line:
                # 添加空行
                self.doc.add_paragraph('')
                i += 1
            # 检查是否为列表项
            elif line.startswith('- ') or line.startswith('* '):
                # 处理列表项
                list_items = [line[2:].strip()]
                i += 1
                while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                    list_items.append(lines[i][2:].strip())
                    i += 1
                
                # 将列表项合并为普通段落
                for item in list_items:
                    self._add_paragraph(f'• {item}', is_title=False)
            # 检查是否为数字列表
            elif re.match(r'^\d+\.\s', line):
                # 处理数字列表
                list_items = [re.sub(r'^\d+\.\s', '', line.strip())]
                i += 1
                while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                    list_items.append(re.sub(r'^\d+\.\s', '', lines[i].strip()))
                    i += 1
                
                # 将列表项合并为普通段落
                for idx, item in enumerate(list_items, 1):
                    self._add_paragraph(f'{idx}. {item}', is_title=False)
            # 普通段落
            else:
                # 收集连续的普通文本
                paragraph_text = line
                i += 1
                while i < len(lines) and not lines[i].strip().startswith(('#', '|', '```', '- ', '* ', '1. ')) and lines[i].strip():
                    paragraph_text += ' ' + lines[i].strip()
                    i += 1
                
                self._add_paragraph(paragraph_text, is_title=False)
        
        # 保存文档
        self.doc.save(output_file)
        print(f"DOCX文档已生成: {output_file}")

if __name__ == "__main__":
    converter = MarkdownToDocxConverter()
    converter.convert(
        r"g:\gitcode\final_assignment\protocol_analyzer\document\用户手册.md",
        r"g:\gitcode\final_assignment\protocol_analyzer\document\用户手册.docx"
    )
