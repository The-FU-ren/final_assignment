#!/usr/bin/env python3
# 将Markdown文档转换为符合要求的DOCX文档

import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def convert_md_to_docx(md_path, docx_path):
    # 读取Markdown文档
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 将Markdown转换为HTML
    html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
    
    # 解析HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # 创建DOCX文档
    doc = Document()
    
    # 设置全局样式
    style = doc.styles['Normal']
    # 设置字体为仿宋
    style.font.name = '仿宋'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    # 设置字体大小为小四（12磅）
    style.font.size = Pt(12)
    # 设置字体颜色为黑色
    style.font.color.rgb = RGBColor(0, 0, 0)
    # 设置单倍行距
    style.paragraph_format.line_spacing = 1.0
    # 设置首行缩进两字符
    style.paragraph_format.first_line_indent = Pt(24)  # 2字符 = 24磅
    
    # 遍历HTML元素
    def process_element(element, parent):
        if element.name == 'h1':
            # 最大标题：黑体加粗三号
            p = parent.add_paragraph()
            p_format = p.paragraph_format
            p_format.line_spacing = 1.0
            p_format.first_line_indent = Pt(0)  # 标题不缩进
            run = p.add_run(element.get_text())
            # 设置字体为黑体
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            # 设置字体大小为三号（16磅）
            run.font.size = Pt(16)
            # 设置字体颜色为黑色
            run.font.color.rgb = RGBColor(0, 0, 0)
            # 设置加粗
            run.font.bold = True
            
        elif element.name == 'h2':
            # 二级标题：黑体加粗小三（15磅）
            p = parent.add_paragraph()
            p_format = p.paragraph_format
            p_format.line_spacing = 1.0
            p_format.first_line_indent = Pt(0)  # 标题不缩进
            run = p.add_run(element.get_text())
            # 设置字体为黑体
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            # 设置字体大小为小三（15磅）
            run.font.size = Pt(15)
            # 设置字体颜色为黑色
            run.font.color.rgb = RGBColor(0, 0, 0)
            # 设置加粗
            run.font.bold = True
            
        elif element.name == 'h3':
            # 三级标题：黑体加粗四号（14磅）
            p = parent.add_paragraph()
            p_format = p.paragraph_format
            p_format.line_spacing = 1.0
            p_format.first_line_indent = Pt(0)  # 标题不缩进
            run = p.add_run(element.get_text())
            # 设置字体为黑体
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            # 设置字体大小为四号（14磅）
            run.font.size = Pt(14)
            # 设置字体颜色为黑色
            run.font.color.rgb = RGBColor(0, 0, 0)
            # 设置加粗
            run.font.bold = True
            
        elif element.name == 'h4':
            # 四级标题：黑体加粗小四（12磅）
            p = parent.add_paragraph()
            p_format = p.paragraph_format
            p_format.line_spacing = 1.0
            p_format.first_line_indent = Pt(0)  # 标题不缩进
            run = p.add_run(element.get_text())
            # 设置字体为黑体
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            # 设置字体大小为小四（12磅）
            run.font.size = Pt(12)
            # 设置字体颜色为黑色
            run.font.color.rgb = RGBColor(0, 0, 0)
            # 设置加粗
            run.font.bold = True
            
        elif element.name == 'p':
            # 正文段落
            p = parent.add_paragraph()
            p_format = p.paragraph_format
            p_format.line_spacing = 1.0
            p_format.first_line_indent = Pt(24)  # 首行缩进两字符
            
            # 处理段落内容
            for child in element.children:
                if child.name == 'strong':
                    # 加粗文本
                    run = p.add_run(child.get_text())
                    run.font.bold = True
                elif child.name == 'em':
                    # 斜体文本
                    run = p.add_run(child.get_text())
                    run.font.italic = True
                elif child.name == 'code':
                    # 代码文本
                    run = p.add_run(child.get_text())
                    run.font.name = 'Courier New'
                else:
                    # 普通文本
                    run = p.add_run(child.get_text())
            
        elif element.name == 'ul' or element.name == 'ol':
            # 列表
            for li in element.find_all('li', recursive=False):
                p = parent.add_paragraph()
                p_format = p.paragraph_format
                p_format.line_spacing = 1.0
                p_format.first_line_indent = Pt(0)  # 列表不缩进
                
                # 添加列表符号
                if element.name == 'ul':
                    p.add_run('• ')
                else:
                    # 有序列表，需要计算序号
                    index = list(element.find_all('li')).index(li) + 1
                    p.add_run(f'{index}. ')
                
                # 处理列表项内容
                for child in li.children:
                    if child.name == 'strong':
                        run = p.add_run(child.get_text())
                        run.font.bold = True
                    elif child.name == 'code':
                        run = p.add_run(child.get_text())
                        run.font.name = 'Courier New'
                    else:
                        run = p.add_run(child.get_text())
        
        elif element.name == 'pre':
            # 代码块
            p = parent.add_paragraph()
            p_format = p.paragraph_format
            p_format.line_spacing = 1.0
            p_format.first_line_indent = Pt(0)  # 代码块不缩进
            p_format.left_indent = Inches(0.5)
            
            code_block = element.find('code')
            if code_block:
                code_text = code_block.get_text()
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
        
        elif element.name == 'table':
            # 表格
            table = parent.add_table(rows=0, cols=0)
            table.style = 'Table Grid'
            
            # 处理表头
            thead = element.find('thead')
            if thead:
                tr = thead.find('tr')
                if tr:
                    headers = tr.find_all('th')
                    # 添加表头行
                    row = table.add_row().cells
                    for i, header in enumerate(headers):
                        # 扩展表格列数（如果需要）
                        if i >= len(row):
                            table.add_column(Inches(2))
                            row = table.rows[0].cells
                        # 设置表头内容
                        cell = row[i]
                        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                        p.add_run(header.get_text())
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        p.paragraph_format.first_line_indent = Pt(0)  # 表格不缩进
            
            # 处理表体
            tbody = element.find('tbody')
            if tbody:
                for tr in tbody.find_all('tr'):
                    row = table.add_row().cells
                    tds = tr.find_all('td')
                    for i, td in enumerate(tds):
                        # 扩展表格列数（如果需要）
                        if i >= len(row):
                            table.add_column(Inches(2))
                            row = table.rows[-1].cells
                        # 设置单元格内容
                        cell = row[i]
                        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                        p.add_run(td.get_text())
                        p.paragraph_format.first_line_indent = Pt(0)  # 表格不缩进
        
        elif element.name == 'div' or element.name == 'body':
            # 容器元素，递归处理子元素
            for child in element.children:
                process_element(child, parent)
        
        elif element.name is None:
            # 文本节点，忽略
            pass
        
        else:
            # 其他元素，递归处理子元素
            for child in element.children:
                process_element(child, parent)
    
    # 处理HTML内容
    process_element(soup, doc)
    
    # 保存DOCX文档
    doc.save(docx_path)
    print(f'转换完成：{docx_path}')

if __name__ == '__main__':
    # 输入和输出文件路径
    md_file = r'g:\gitcode\final_assignment\protocol_analyzer\document\软件设计文档.md'
    docx_file = r'g:\gitcode\final_assignment\protocol_analyzer\document\软件设计文档.docx'
    
    # 执行转换
    convert_md_to_docx(md_file, docx_file)